from __future__ import annotations

from pathlib import Path

from .base import BaseExtractor, ExtractionResult, ExtractorError


class DocxExtractor(BaseExtractor):
    name = "docx"
    extensions = (".docx",)

    def extract(self, path: Path) -> ExtractionResult:
        try:
            from docx import Document
        except ImportError as e:
            raise ExtractorError(f"python-docx not installed: {e}")

        try:
            doc = Document(str(path))
        except Exception as e:
            raise ExtractorError(f"DOCX open failed: {e}")

        parts: list[str] = []

        for para in doc.paragraphs:
            if para.text:
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text]
                if row_cells:
                    parts.append("\t".join(row_cells))

        return ExtractionResult(text="\n".join(parts))
